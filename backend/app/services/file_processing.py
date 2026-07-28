"""文件处理微服务 (层级 4): 文档解析 / 语义分块 / 元数据注入 / 向量化入库 (RAG).

设计要点对应（来自生产级 RAG 规范）：
  (1) 数据治理：解析后做基础清洗（去页码水印等噪声）。
  (2) 语义分块：按 doc_type 匹配分块策略（FAQ / 制度 / 技术手册 / 长文），拒绝一刀切。
      解析：PyMuPDF(fitz) 解析 PDF，python-docx 解析 Word，openpyxl 解析 Excel，
           markdown-it-py 规范化 Markdown 后交由 MarkdownHeaderTextSplitter 按标题切片。
  (4) 元数据与溯源：每个分块注入 meta（来源/页码/章节/文档类型/权限/标签/更新时间）。
  (5) 增量更新：每次上传只索引当前文档，不重建全库（由调用方触发）。
向量：bge-small-zh-v1.5；重排：bge-reranker-v2-m3（见 app/agent/rag.py）。
"""
from __future__ import annotations

import io
import logging
import re
from datetime import datetime

from app.agent.rag import embed_text, embed_texts
from app.config import settings
from app.core.storage import save_object
from app.db.models import Document, DocumentChunk

logger = logging.getLogger("file_processing")

# --------------------------------------------------------------------------- #
# 分块策略：按文档类型匹配（token 估算 ≈ 1.6 中文字符/token）
# --------------------------------------------------------------------------- #
CHUNK_STRATEGY = {
    "faq":        {"splitter": "faq",      "chunk_size": 320,  "chunk_overlap": 0},    # 100-200 tokens：单条问答完整保留
    "regulation": {"splitter": "recursive", "chunk_size": 600,  "chunk_overlap": 120},  # ~375 tokens：单条/单节一款，便于精准命中
    "manual":     {"splitter": "markdown",  "chunk_size": 700,  "chunk_overlap": 100},  # ~440 tokens：按标题层级
    "book":       {"splitter": "recursive", "chunk_size": 1200, "chunk_overlap": 150},  # ~750 tokens：段落+章节
    "general":    {"splitter": "recursive", "chunk_size": 500,  "chunk_overlap": 80},
}

DOC_TYPES = set(CHUNK_STRATEGY.keys())


# ---------- 解析（返回按页切分的文本，便于溯源页码） ----------

def extract_pages(filename: str, data: bytes) -> list[tuple[int | None, str]]:
    """返回 [(page_no, text), ...]；非分页文档 page_no=None。"""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return _extract_pdf_pages(data)
    if lower.endswith(".docx"):
        return [(None, _extract_docx(data))]
    if lower.endswith((".xlsx", ".xlsm")):
        return [(None, _extract_xlsx(data))]
    if lower.endswith((".txt", ".md", ".markdown")):
        return [(None, data.decode("utf-8", errors="ignore"))]
    return [(None, f"[不支持解析的文件类型: {filename}]")]


def _extract_pdf_pages(data: bytes) -> list[tuple[int, str]]:
    try:
        import fitz  # PyMuPDF
    except Exception:
        return []
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        try:
            return [(i + 1, page.get_text("text")) for i, page in enumerate(doc)]
        finally:
            doc.close()
    except Exception:
        return []


def _extract_docx(data: bytes) -> str:
    try:
        import docx
    except Exception:
        return ""
    try:
        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text for c in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_xlsx(data: bytes) -> str:
    try:
        import openpyxl
    except Exception:
        return ""
    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        out = []
        for ws in wb.worksheets:
            out.append(f"# {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = ["" if c is None else str(c) for c in row]
                if any(cells):
                    out.append("\t".join(cells))
        return "\n".join(out)
    except Exception:
        return ""


def _normalize_markdown(text: str) -> str:
    """用 markdown-it-py 将 Markdown 规范化为保留标题语法的纯文本。"""
    try:
        from markdown_it import MarkdownIt
    except Exception:
        return text
    try:
        tokens = MarkdownIt("commonmark").parse(text)
        lines: list[str] = []
        for i, tok in enumerate(tokens):
            if tok.type == "heading_open":
                level = int(tok.tag[1])
                inline = tokens[i + 1] if i + 1 < len(tokens) else None
                content = inline.content if inline and inline.type == "inline" else ""
                lines.append("#" * level + " " + content)
            elif tok.type == "paragraph_open":
                inline = tokens[i + 1] if i + 1 < len(tokens) else None
                if inline and inline.type == "inline" and inline.content:
                    lines.append(inline.content)
            elif tok.type == "list_item_open":
                inline = tokens[i + 1] if i + 1 < len(tokens) else None
                if inline and inline.type == "inline" and inline.content:
                    lines.append("- " + inline.content)
        return "\n".join(lines)
    except Exception:
        return text


# ---------- 数据治理：基础清洗 ----------

def _clean_text(text: str) -> str:
    """去除常见噪声：空白归一、页码（“第 N 页”）、纯装饰线、空行。"""
    if not text:
        return ""
    lines = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if re.fullmatch(r"第\s*\d+\s*页", line):  # 页码
            continue
        if re.fullmatch(r"[-=*_•·]{3,}", line):  # 装饰线
            continue
        lines.append(line)
    return "\n".join(lines)


# ---------- 切片（返回 [{content, section}]） ----------

# 制度类文档的章节标题边界（如「第九章 安全管理制度」「第一章 总则」「三、考勤管理」）。
# docx/纯文本没有 markdown 标题层级，用该正则把整篇制度按章节切分为独立 chunk，
# 避免「安全管理制度」与「离职交接」「附则」被切进同一块，导致答非所问。
_SECTION_PAT = re.compile(
    r"^\s*(?:"
    r"(?:第[一二三四五六七八九十百零\d]+[章条节篇编部分]|[CHS]\.?\d+)"
    r"[\s、.．:-]*[^\s\d][^\n]{0,30}"  # 标题文字（如「安全管理制度」）
    r"|"
    r"\d+[\s、.．:-]+[^\s\d][^\n]{0,20}"  # 形如「1、入职管理」「3. 考勤」
    r")\s*$"
)

# 仅当一行「看起来像章节标题」时才切分：标题文字不能太短（>=2 字），避免把正文序号行误切
def _looks_like_section(line: str) -> bool:
    m = _SECTION_PAT.match(line)
    if not m:
        return False
    title_text = line.strip()
    # 标题过短（如只剩「第九章」「第三条」且无实质名称）时，并入下文不单独切片
    body = re.sub(r"^(?:第[一二三四五六七八九十百零\d]+[章条节篇编部分]|[CHS]\.?\d+|[0-9]+)[\s、.．:-]*", "", title_text)
    return len(body) >= 2


def _split_by_sections(text: str, strategy: dict) -> list[dict]:
    """按章节标题边界切分制度类文档，使每个章节成为独立 chunk。

    章节内若仍超 chunk_size，再按递归切分器细切（带章节名作为 section 上下文）。
    """
    lines = text.splitlines()
    blocks: list[list[str]] = []
    cur: list[str] = []
    cur_section = ""

    def flush() -> None:
        nonlocal cur, cur_section
        if cur:
            blocks.append((cur_section, "\n".join(cur).strip()))
        cur, cur_section = [], ""

    for line in lines:
        if _looks_like_section(line):
            flush()
            cur_section = line.strip()
            cur = [line]
        else:
            if not cur and not cur_section:
                # 文档开头的无标题导语（如前言），归入首个临时块
                cur_section = ""
            cur.append(line)
    flush()

    out: list[dict] = []
    for section, content in blocks:
        if not content:
            continue
        if len(content) > strategy["chunk_size"]:
            for sub in _split_recursive(content, strategy):
                out.append({"content": sub["content"], "section": section})
        else:
            out.append({"content": content, "section": section})
    # 兜底：若整篇未识别出任何章节（纯段落流），回退递归切片
    if not out:
        out = _split_recursive(text, strategy)
    return out


def split_document(filename: str, text: str, doc_type: str = "general") -> list[dict]:
    lower = filename.lower()
    strategy = CHUNK_STRATEGY.get(doc_type, CHUNK_STRATEGY["general"])
    kind = strategy["splitter"]

    if lower.endswith((".md", ".markdown")):
        md_text = _normalize_markdown(text)
        return _split_markdown(md_text)
    if kind == "faq":
        chunks = _split_faq(text)
        if len(chunks) > 1:
            return [{"content": c, "section": ""} for c in chunks]
        return _split_recursive(text, strategy)  # 无问答结构则回退
    if kind == "markdown":
        return _split_markdown(_normalize_markdown(text))
    # 制度类 / 通用类文档：按章节边界切分（doc_type 决定，而非 splitter 字段）。
    # 保证「第X章/第X条」各自独立成块，检索精准（如「安全管理制度」不会混入
    # 「离职交接」「附则」等无关章节）。若未识别出章节结构，则回退到递归切片。
    if doc_type in ("regulation", "general"):
        return _split_by_sections(text, strategy)
    return _split_recursive(text, strategy)


def _split_markdown(text: str) -> list[dict]:
    try:
        from langchain_text_splitters import MarkdownHeaderTextSplitter
    except Exception:
        return [{"content": t, "section": ""} for t in _split_recursive(text, CHUNK_STRATEGY["manual"])]
    headers_to_split_on = [
        ("#", "H1"), ("##", "H2"), ("###", "H3"), ("####", "H4"),
    ]
    splitter = MarkdownHeaderTextSplitter(headers_to_split_on=headers_to_split_on)
    try:
        splits = splitter.split_text(text)
    except Exception:
        return [{"content": t, "section": ""} for t in _split_recursive(text, CHUNK_STRATEGY["manual"])]
    out = []
    for s in splits:
        section = " > ".join(f"{k}:{v}" for k, v in (s.metadata or {}).items())
        content = (s.page_content).strip()
        if content:
            out.append({"content": content, "section": section})
    if not out:
        out = [{"content": text[:400], "section": ""}]
    return out


def _split_faq(text: str) -> list[str]:
    """保留单条问答完整性，不跨问题切割。"""
    q_pat = re.compile(r"^\s*(问|Q|q)[\s:：.)]")
    a_pat = re.compile(r"^\s*(答|A|a)[\s:：.)]")
    chunks: list[str] = []
    cur_q: str | None = None
    buf: list[str] = []

    def flush():
        nonlocal cur_q, buf
        if cur_q is not None or buf:
            chunk = (cur_q + "\n" + "\n".join(buf)).strip() if cur_q else "\n".join(buf).strip()
            if chunk:
                chunks.append(chunk)
        cur_q, buf = None, []

    for line in text.splitlines():
        line = line.rstrip()
        if q_pat.match(line):
            flush()
            cur_q = line
        elif a_pat.match(line):
            if cur_q is None:
                cur_q = line
            else:
                buf.append(line)
        else:
            if cur_q is not None or buf:
                buf.append(line)
    flush()
    return chunks


def _split_recursive(text: str, strategy: dict) -> list[dict]:
    text = text.strip()
    if not text:
        return []
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=strategy["chunk_size"],
            chunk_overlap=strategy["chunk_overlap"],
            separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
        )
        return [{"content": c, "section": ""} for c in (splitter.split_text(text) or [text[:strategy["chunk_size"]]])]
    except Exception:
        return [{"content": c, "section": ""} for c in _fallback_chunk(text, strategy["chunk_size"])]


def _fallback_chunk(text: str, size: int) -> list[str]:
    paras = [p for p in text.split("\n") if p.strip()]
    chunks, buf = [], ""
    for p in paras:
        if len(buf) + len(p) > size:
            if buf:
                chunks.append(buf)
            buf = p
        else:
            buf += "\n" + p if buf else p
    if buf:
        chunks.append(buf)
    return chunks or [text[:size]]


# ---------- 入库（增量，带元数据） ----------

async def ingest_file(
    session,
    data: bytes,
    filename: str,
    department: str,
    user_id: str,
    *,
    doc_type: str = "general",
    trust_level: str = "internal",
    tags: list | None = None,
    source: str = "",
) -> Document:
    doc_type = doc_type if doc_type in DOC_TYPES else "general"
    tags = tags or []
    now_iso = datetime.now().isoformat(timespec="seconds")

    text_all = _clean_text("\n".join(t for _, t in extract_pages(filename, data)))
    stored = await save_object(data, filename, content_type="application/octet-stream")

    doc = Document(
        title=filename,
        department=department,
        file_path=stored["object_key"],
        doc_type=doc_type,
        trust_level=trust_level,
        source=source or filename,
        tags=tags,
        version=1,
        doc_status="active",
        created_by=user_id,
    )
    session.add(doc)
    await session.flush()

    # 先收集所有分块（含元数据），再批量向量化，避免逐 chunk 串行 HTTP 请求
    parts_meta: list[tuple[int | None, dict]] = []
    for page_no, raw in extract_pages(filename, data):
        cleaned = _clean_text(raw)
        if not cleaned:
            continue
        for part in split_document(filename, cleaned, doc_type):
            parts_meta.append((page_no, part))

    embeddings = await embed_texts([p["content"] for _, p in parts_meta])

    chunks: list[DocumentChunk] = []
    for idx, ((page_no, part), emb) in enumerate(zip(parts_meta, embeddings)):
        chunks.append(DocumentChunk(
            document_id=doc.id,
            department=department,
            content=part["content"],
            chunk_index=idx,
            embedding=emb,
            meta={
                "source": doc.source or filename,
                "title": filename,
                "page": page_no,
                "section": part["section"],
                "doc_type": doc_type,
                "permission": department,
                "tags": tags,
                "updated_at": now_iso,
            },
        ))
    session.add_all(chunks)
    doc.chunk_count = len(chunks)
    await session.commit()
    await session.refresh(doc)
    return doc
